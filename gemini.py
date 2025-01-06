import os
import requests
import json
from docx import Document
from docx.shared import RGBColor
from urllib.parse import urlparse
import google.generativeai as genai

import re

# Set your Gemini API key
os.environ["GEMINI_API_KEY"] = "your-api-key"
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

base_url = "https://clinicaltrials.gov/api/v2/studies"


def generate_question_and_answer(gene, document_context, max_retries):
    question = f"Based on the following clinical trial information, would a patient with a {gene} gene mutation be eligible for the clinical trial? {document_context}"
    retry_count = 0
    model = genai.GenerativeModel("gemini-1.5-pro")
    while retry_count < max_retries:
        try:
            response = model.generate_content(question)
            answer = response.text
            question1 = f"Based on the following clinical trial information, would a patient with a {gene} gene mutation be eligible for the clinical trial?"
            return question1, answer
        except Exception:
            retry_count += 1

    return None, None


def filter_unwanted_sections(data):
    filtered_data = []
    wanted_modules = ["eligibilityCriteria", "healthyVolunteers", "sex", "minimumAge", "studyPopulation"]
    inside_stdages = False

    for line in data.splitlines():

        if any(module in line for module in wanted_modules):
            skip = False
        else:
            skip = True
        if '"stdAges": [' in line:
            inside_stdages = True
            filtered_data.append(line.strip())  # Add the opening line of the stdAges array
            continue
        elif inside_stdages and ']' in line:
            inside_stdages = False
            filtered_data.append(line.strip())  # Add the closing bracket of the stdAges array
            continue
        if inside_stdages:
            # Capture elements inside the stdAges array, assuming a single line array for simplicity
            if '"' in line:
                filtered_data.append(line.strip())
            continue
        if skip:
            continue

        # Collect lines that are not within unwanted sections
        filtered_data.append(line)
    return "\n".join(filtered_data)


# Function to extract nctId from URL
def extract_nct_id_from_url(url):
    path = urlparse(url).path
    nct_id = path.split('/')[-1]
    return nct_id

# Function to fetch data for a list of URLs and create documents
def fetch_data_and_create_documents_from_links(links, folder_name, file_prefix, gene, gene_short):
    data_list = []

    for link in links:
        nctId = extract_nct_id_from_url(link)
        response = requests.get(f"{base_url}/{nctId}")

        if response.status_code == 200:
            study = response.json()  # Parse JSON response
            title = study['protocolSection']['identificationModule'].get('briefTitle', 'Unknown')

            # Collect all available information
            all_info = json.dumps(study, indent=2)
            filtered_info = filter_unwanted_sections(all_info)

            # Append the data to the list as a dictionary
            data_list.append({
                "Link": link,
                "AllInfo": filtered_info,
                "Title": title
            })
        else:
            print(f"Failed to fetch data for {nctId}. Status code:", response.status_code)

    # Create a folder to store documents
    folder_path = f"C:/path/to/your/folder/{folder_name}"
    os.makedirs(folder_path, exist_ok=True)

    # Create documents with 1 trial per document
    num_docs = len(data_list)  # Number of documents equals the number of trials
    doc_count = 0

    for i in range(num_docs):

        doc = Document()

        # Add a title to the document
        doc.add_heading(f'Clinical Trials Data {gene_short} - Document {i + 1}', level=1)

        # Add data to the document
        entry = data_list[i]
        doc.add_heading(f"{entry['Title']}")
        doc.add_heading(f"Clinical Trial: {entry['Link']}", level=2)
        doc.add_paragraph(filter_unwanted_sections(entry['AllInfo']))

        # Generate question and answer

        document_context = filter_unwanted_sections(entry['AllInfo'])

        question, answer = generate_question_and_answer(gene, document_context, 10000)


        p_question = doc.add_paragraph()
        p_question.add_run("Question:\n").font.color.rgb = RGBColor(255, 0, 0)  # Red color
        if question is not None:
            p_question.add_run(question).font.color.rgb = RGBColor(255, 0, 0)
        else:
            p_question.add_run("No question available").font.color.rgb = RGBColor(255, 0, 0)
            print(f'Clinical Trials Data {gene_short} - Document {i + 1} could not get an answer')

        p_answer = doc.add_paragraph()
        p_answer.add_run("Answer:\n").font.color.rgb = RGBColor(255, 0, 0)  # Red color
        if answer is not None:
            p_answer.add_run(answer).font.color.rgb = RGBColor(255, 0, 0)
        else:
            p_answer.add_run("No answer available").font.color.rgb = RGBColor(255, 0, 0)
            print(f'Clinical Trials Data {gene_short} - Document {i + 1} could not get an answer')

        doc.add_paragraph("")  # Empty paragraph for spacing

        # Save the document
        doc_file_path = os.path.join(folder_path, f"{file_prefix}_data_{gene_short}{i + 1}.docx")
        doc.save(doc_file_path)
        doc_count += 1

    print("All documents saved successfully.")


# List of Clinical Trials URLs
links_ALK = [
    "https://clinicaltrials.gov/study/NCT01838577",
    "https://clinicaltrials.gov/study/NCT01752400",
    "https://clinicaltrials.gov/study/NCT01574300",
    "https://clinicaltrials.gov/study/NCT01904916",
    "https://clinicaltrials.gov/study/NCT01994057",
    "https://clinicaltrials.gov/study/NCT01579994",
    "https://clinicaltrials.gov/study/NCT01742286",
    "https://clinicaltrials.gov/study/NCT01100840",
    "https://clinicaltrials.gov/study/NCT01979263",
    "https://clinicaltrials.gov/study/NCT01629719",
    "https://clinicaltrials.gov/study/NCT01306656",
    "https://clinicaltrials.gov/study/NCT01145937",
    "https://clinicaltrials.gov/study/NCT02040870",
    "https://clinicaltrials.gov/study/NCT01822496",
    "https://clinicaltrials.gov/study/NCT01662635",
    "https://clinicaltrials.gov/study/NCT01998126",
    "https://clinicaltrials.gov/study/NCT01930474",
    "https://clinicaltrials.gov/study/NCT01852825",
    "https://clinicaltrials.gov/study/NCT00479973",
    "https://clinicaltrials.gov/study/NCT01596374",
    "https://clinicaltrials.gov/study/NCT01712217",
    "https://clinicaltrials.gov/study/NCT02041468",
    "https://clinicaltrials.gov/study/NCT02106169",
    "https://clinicaltrials.gov/study/NCT02085135",
    "https://clinicaltrials.gov/study/NCT01625234",
    "https://clinicaltrials.gov/study/NCT02069535",
    "https://clinicaltrials.gov/study/NCT01829503",
    "https://clinicaltrials.gov/study/NCT02024087",
    "https://clinicaltrials.gov/study/NCT01829217",
    "https://clinicaltrials.gov/study/NCT01999972",
    "https://clinicaltrials.gov/study/NCT02171286",

]
links_BRAF = [
    "https://clinicaltrials.gov/study/NCT01838577",
    "https://clinicaltrials.gov/study/NCT02015117",
    "https://clinicaltrials.gov/study/NCT01693419",
    "https://clinicaltrials.gov/study/NCT01260415",
    "https://clinicaltrials.gov/study/NCT01750918",
    "https://clinicaltrials.gov/study/NCT01740648",
    "https://clinicaltrials.gov/study/NCT01449058",
    "https://clinicaltrials.gov/study/NCT01358812",
    "https://clinicaltrials.gov/study/NCT01704703",
    "https://clinicaltrials.gov/study/NCT01124669",
    "https://clinicaltrials.gov/study/NCT01282502",
    "https://clinicaltrials.gov/study/NCT01110785",
    "https://clinicaltrials.gov/study/NCT01787500",
    "https://clinicaltrials.gov/study/NCT01719380",
    "https://clinicaltrials.gov/study/NCT00326495",
    "https://clinicaltrials.gov/study/NCT01758575",
    "https://clinicaltrials.gov/study/NCT01306045",
    "https://clinicaltrials.gov/study/NCT00849407",
    "https://clinicaltrials.gov/study/NCT01100840",
    "https://clinicaltrials.gov/study/NCT02038348",
    "https://clinicaltrials.gov/study/NCT01738451",
    "https://clinicaltrials.gov/study/NCT01657591",
    "https://clinicaltrials.gov/study/NCT00991991",
    "https://clinicaltrials.gov/study/NCT01907802",
    "https://clinicaltrials.gov/study/NCT01512251",
    "https://clinicaltrials.gov/study/NCT01089101",
    "https://clinicaltrials.gov/study/NCT01876511",
    "https://clinicaltrials.gov/study/NCT02042040",
    "https://clinicaltrials.gov/study/NCT01954043",
    "https://clinicaltrials.gov/study/NCT01802645",
    "https://clinicaltrials.gov/study/NCT01659151",
    "https://clinicaltrials.gov/study/NCT01377025",
    "https://clinicaltrials.gov/study/NCT01667419",
    "https://clinicaltrials.gov/study/NCT01640444",
    "https://clinicaltrials.gov/study/NCT01959633",
    "https://clinicaltrials.gov/study/NCT01878396",
    "https://clinicaltrials.gov/study/NCT02142218",
    "https://clinicaltrials.gov/study/NCT01596140",
    "https://clinicaltrials.gov/study/NCT01682083",
    "https://clinicaltrials.gov/study/NCT01711632",
    "https://clinicaltrials.gov/study/NCT01688232",
    "https://clinicaltrials.gov/study/NCT02145910",
    "https://clinicaltrials.gov/study/NCT02130466",
    "https://clinicaltrials.gov/study/NCT01894672",
    "https://clinicaltrials.gov/study/NCT01713972",
    "https://clinicaltrials.gov/study/NCT01791309",
    "https://clinicaltrials.gov/study/NCT01841463",
    "https://clinicaltrials.gov/study/NCT01586195",
    "https://clinicaltrials.gov/study/NCT02083354",
    "https://clinicaltrials.gov/study/NCT02034110",
    "https://clinicaltrials.gov/study/NCT02097225",
    "https://clinicaltrials.gov/study/NCT01585415",
    "https://clinicaltrials.gov/study/NCT01754376",
    "https://clinicaltrials.gov/study/NCT01543698",
    "https://clinicaltrials.gov/study/NCT01972347",
    "https://clinicaltrials.gov/study/NCT01781026",
    "https://clinicaltrials.gov/study/NCT02171286",

]
links_EGFR = [
    "https://clinicaltrials.gov/study/NCT01838577",
    "https://clinicaltrials.gov/study/NCT01380795",
    "https://clinicaltrials.gov/study/NCT00842257",
    "https://clinicaltrials.gov/study/NCT01294826",
    "https://clinicaltrials.gov/study/NCT01384994",
    "https://clinicaltrials.gov/study/NCT01693419",
    "https://clinicaltrials.gov/study/NCT01260415",
    "https://clinicaltrials.gov/study/NCT01750918",
    "https://clinicaltrials.gov/study/NCT01358812",
    "https://clinicaltrials.gov/study/NCT01740804",
    "https://clinicaltrials.gov/study/NCT01542437",
    "https://clinicaltrials.gov/study/NCT01124669",
    "https://clinicaltrials.gov/study/NCT01574300",
    "https://clinicaltrials.gov/study/NCT01596790",
    "https://clinicaltrials.gov/study/NCT01719380",
    "https://clinicaltrials.gov/study/NCT01608841",
    "https://clinicaltrials.gov/study/NCT02136550",
    "https://clinicaltrials.gov/study/NCT01620190",
    "https://clinicaltrials.gov/study/NCT01497626",
    "https://clinicaltrials.gov/study/NCT01646450",
    "https://clinicaltrials.gov/study/NCT01647711",
    "https://clinicaltrials.gov/study/NCT02148380",
    "https://clinicaltrials.gov/study/NCT01285375",
    "https://clinicaltrials.gov/study/NCT01741727",
    "https://clinicaltrials.gov/study/NCT01420874",
    "https://clinicaltrials.gov/study/NCT01273610",
    "https://clinicaltrials.gov/study/NCT01454102",
    "https://clinicaltrials.gov/study/NCT02070679",
    "https://clinicaltrials.gov/study/NCT01787006",
    "https://clinicaltrials.gov/study/NCT01532089",
    "https://clinicaltrials.gov/study/NCT01305772",
    "https://clinicaltrials.gov/study/NCT01767974",
    "https://clinicaltrials.gov/study/NCT01697163",
    "https://clinicaltrials.gov/study/NCT02141672",
    "https://clinicaltrials.gov/study/NCT01931306",
    "https://clinicaltrials.gov/study/NCT02147990",
    "https://clinicaltrials.gov/study/NCT02025114",
    "https://clinicaltrials.gov/study/NCT00863122",
    "https://clinicaltrials.gov/study/NCT01719536",
    "https://clinicaltrials.gov/study/NCT01996098",
    "https://clinicaltrials.gov/study/NCT00326495",
    "https://clinicaltrials.gov/study/NCT02091960",
    "https://clinicaltrials.gov/study/NCT01858389",
    "https://clinicaltrials.gov/study/NCT00353717",
    "https://clinicaltrials.gov/study/NCT01723774",
    "https://clinicaltrials.gov/study/NCT01109095",
    "https://clinicaltrials.gov/study/NCT01605266",
    "https://clinicaltrials.gov/study/NCT01228045",
    "https://clinicaltrials.gov/study/NCT01717807",
    "https://clinicaltrials.gov/study/NCT01391260",
    "https://clinicaltrials.gov/study/NCT02001896",
    "https://clinicaltrials.gov/study/NCT00569296",
    "https://clinicaltrials.gov/study/NCT01953913",
    "https://clinicaltrials.gov/study/NCT00984425",
    "https://clinicaltrials.gov/study/NCT01806649",
    "https://clinicaltrials.gov/study/NCT01394120",
    "https://clinicaltrials.gov/study/NCT01758575",
    "https://clinicaltrials.gov/study/NCT01904916",
    "https://clinicaltrials.gov/study/NCT01306045",
    "https://clinicaltrials.gov/study/NCT01665417",
    "https://clinicaltrials.gov/study/NCT01048918",
    "https://clinicaltrials.gov/study/NCT01833572",
    "https://clinicaltrials.gov/study/NCT02063906",
    "https://clinicaltrials.gov/study/NCT01943786",
    "https://clinicaltrials.gov/study/NCT01494662",
    "https://clinicaltrials.gov/study/NCT02107703",
    "https://clinicaltrials.gov/study/NCT02140333",
    "https://clinicaltrials.gov/study/NCT02122172",
    "https://clinicaltrials.gov/study/NCT01407822",
    "https://clinicaltrials.gov/study/NCT00902044",
    "https://clinicaltrials.gov/study/NCT01534585",
    "https://clinicaltrials.gov/study/NCT01829178",
    "https://clinicaltrials.gov/study/NCT01951469",
    "https://clinicaltrials.gov/study/NCT01763307",
    "https://clinicaltrials.gov/study/NCT01692418",
    "https://clinicaltrials.gov/study/NCT01779050",
    "https://clinicaltrials.gov/study/NCT01292356",
    "https://clinicaltrials.gov/study/NCT02049957",
    "https://clinicaltrials.gov/study/NCT01405079",
    "https://clinicaltrials.gov/study/NCT02036359",
    "https://clinicaltrials.gov/study/NCT01965275",
    "https://clinicaltrials.gov/study/NCT01513174",
    "https://clinicaltrials.gov/study/NCT01998789",
    "https://clinicaltrials.gov/study/NCT00809237",
    "https://clinicaltrials.gov/study/NCT01348412",
    "https://clinicaltrials.gov/study/NCT01376505",
    "https://clinicaltrials.gov/study/NCT02145637",
    "https://clinicaltrials.gov/study/NCT01580865",
    "https://clinicaltrials.gov/study/NCT01967095",
    "https://clinicaltrials.gov/study/NCT01941654",
    "https://clinicaltrials.gov/study/NCT01848756",
    "https://clinicaltrials.gov/study/NCT01784549",
    "https://clinicaltrials.gov/study/NCT01465802",
    "https://clinicaltrials.gov/study/NCT01861223",
    "https://clinicaltrials.gov/study/NCT01873833",
    "https://clinicaltrials.gov/study/NCT01854034",
    "https://clinicaltrials.gov/study/NCT01728233",
    "https://clinicaltrials.gov/study/NCT01730833",
    "https://clinicaltrials.gov/study/NCT00940316",
    "https://clinicaltrials.gov/study/NCT01785420",
    "https://clinicaltrials.gov/study/NCT00452075",
    "https://clinicaltrials.gov/study/NCT01730118",
    "https://clinicaltrials.gov/study/NCT01393080",
    "https://clinicaltrials.gov/study/NCT01627379",
    "https://clinicaltrials.gov/study/NCT01805362",
    "https://clinicaltrials.gov/study/NCT02017171",
    "https://clinicaltrials.gov/study/NCT00970502",
    "https://clinicaltrials.gov/study/NCT00899405",
    "https://clinicaltrials.gov/study/NCT00889954",
    "https://clinicaltrials.gov/study/NCT01994057",
    "https://clinicaltrials.gov/study/NCT02047903",
    "https://clinicaltrials.gov/study/NCT01892527",
    "https://clinicaltrials.gov/study/NCT01000428",
    "https://clinicaltrials.gov/study/NCT02125240",
    "https://clinicaltrials.gov/study/NCT01993784",
    "https://clinicaltrials.gov/study/NCT02013089",
    "https://clinicaltrials.gov/study/NCT01688713",
    "https://clinicaltrials.gov/study/NCT02117167",
    "https://clinicaltrials.gov/study/NCT00601913",
    "https://clinicaltrials.gov/study/NCT01131429",
    "https://clinicaltrials.gov/study/NCT01724801",
    "https://clinicaltrials.gov/study/NCT01523340",
    "https://clinicaltrials.gov/study/NCT00950417",
    "https://clinicaltrials.gov/study/NCT01360931",
    "https://clinicaltrials.gov/study/NCT01976169",
    "https://clinicaltrials.gov/study/NCT02069730",
    "https://clinicaltrials.gov/study/NCT00939523",
    "https://clinicaltrials.gov/study/NCT01526473",
    "https://clinicaltrials.gov/study/NCT01937689",
    "https://clinicaltrials.gov/study/NCT01922921",
    "https://clinicaltrials.gov/study/NCT01602406",
    "https://clinicaltrials.gov/study/NCT01757327",
    "https://clinicaltrials.gov/study/NCT02102438",
    "https://clinicaltrials.gov/study/NCT01989780",
    "https://clinicaltrials.gov/study/NCT01973660",
    "https://clinicaltrials.gov/study/NCT01912963",
    "https://clinicaltrials.gov/study/NCT01100840",
    "https://clinicaltrials.gov/study/NCT01822496",
    "https://clinicaltrials.gov/study/NCT01998126",
    "https://clinicaltrials.gov/study/NCT01930474",
    "https://clinicaltrials.gov/study/NCT02041468",
    "https://clinicaltrials.gov/study/NCT01791309",
    "https://clinicaltrials.gov/study/NCT01649284",
    "https://clinicaltrials.gov/study/NCT01874171",
    "https://clinicaltrials.gov/study/NCT02159495",
    "https://clinicaltrials.gov/study/NCT01829217",
    "https://clinicaltrials.gov/study/NCT02171286",
    "https://clinicaltrials.gov/study/NCT01816035",
    "https://clinicaltrials.gov/study/NCT01957332",

]
links_ERBB2 = [
    "https://clinicaltrials.gov/study/NCT01542437",
    "https://clinicaltrials.gov/study/NCT01273610",
    "https://clinicaltrials.gov/study/NCT00863122",
    "https://clinicaltrials.gov/study/NCT02091960",
    "https://clinicaltrials.gov/study/NCT01723774",
    "https://clinicaltrials.gov/study/NCT01109095",
    "https://clinicaltrials.gov/study/NCT01228045",
    "https://clinicaltrials.gov/study/NCT00984425",
    "https://clinicaltrials.gov/study/NCT01306045",
    "https://clinicaltrials.gov/study/NCT02063906",
    "https://clinicaltrials.gov/study/NCT01494662",
    "https://clinicaltrials.gov/study/NCT02122172",
    "https://clinicaltrials.gov/study/NCT00902044",
    "https://clinicaltrials.gov/study/NCT01779050",
    "https://clinicaltrials.gov/study/NCT02049957",
    "https://clinicaltrials.gov/study/NCT01376505",
    "https://clinicaltrials.gov/study/NCT01848756",
    "https://clinicaltrials.gov/study/NCT01465802",
    "https://clinicaltrials.gov/study/NCT01873833",
    "https://clinicaltrials.gov/study/NCT01730833",
    "https://clinicaltrials.gov/study/NCT01785420",
    "https://clinicaltrials.gov/study/NCT01730118",
    "https://clinicaltrials.gov/study/NCT00889954",
    "https://clinicaltrials.gov/study/NCT01976169",
    "https://clinicaltrials.gov/study/NCT01248897",
    "https://clinicaltrials.gov/study/NCT00939523",
    "https://clinicaltrials.gov/study/NCT01526473",
    "https://clinicaltrials.gov/study/NCT01937689",
    "https://clinicaltrials.gov/study/NCT01861054",
    "https://clinicaltrials.gov/study/NCT01922921",
    "https://clinicaltrials.gov/study/NCT02073916",
    "https://clinicaltrials.gov/study/NCT01456455",
    "https://clinicaltrials.gov/study/NCT00912275",
    "https://clinicaltrials.gov/study/NCT01602406",
    "https://clinicaltrials.gov/study/NCT01354522",
    "https://clinicaltrials.gov/study/NCT00896909",
    "https://clinicaltrials.gov/study/NCT00250874",
    "https://clinicaltrials.gov/study/NCT01662128",
    "https://clinicaltrials.gov/study/NCT02066532",
    "https://clinicaltrials.gov/study/NCT01593020",
    "https://clinicaltrials.gov/study/NCT01935843",
    "https://clinicaltrials.gov/study/NCT01340430",
    "https://clinicaltrials.gov/study/NCT01612546",
    "https://clinicaltrials.gov/study/NCT01757327",
    "https://clinicaltrials.gov/study/NCT02156648",
    "https://clinicaltrials.gov/study/NCT00411788",
    "https://clinicaltrials.gov/study/NCT01344837",
    "https://clinicaltrials.gov/study/NCT01160094",
    "https://clinicaltrials.gov/study/NCT00842998",
    "https://clinicaltrials.gov/study/NCT02102438",
    "https://clinicaltrials.gov/study/NCT01989780",
    "https://clinicaltrials.gov/study/NCT00896727",
    "https://clinicaltrials.gov/study/NCT02000596",
    "https://clinicaltrials.gov/study/NCT01325207",
    "https://clinicaltrials.gov/study/NCT01641406",
    "https://clinicaltrials.gov/study/NCT01973660",
    "https://clinicaltrials.gov/study/NCT01924351",
    "https://clinicaltrials.gov/study/NCT01856036",
    "https://clinicaltrials.gov/study/NCT01912963",
    "https://clinicaltrials.gov/study/NCT01100840",
    "https://clinicaltrials.gov/study/NCT01816035",
    "https://clinicaltrials.gov/study/NCT01957332",

]
links_KIT = [
    "https://clinicaltrials.gov/study/NCT01306045",
    "https://clinicaltrials.gov/study/NCT02013089",
    "https://clinicaltrials.gov/study/NCT00318266",
    "https://clinicaltrials.gov/study/NCT01994213",
    "https://clinicaltrials.gov/study/NCT00653094",
    "https://clinicaltrials.gov/study/NCT02072031",
    "https://clinicaltrials.gov/study/NCT01762293",
    "https://clinicaltrials.gov/study/NCT02093520",
    "https://clinicaltrials.gov/study/NCT01525550",
    "https://clinicaltrials.gov/study/NCT00646633",
    "https://clinicaltrials.gov/study/NCT01931007",
    "https://clinicaltrials.gov/study/NCT02069730",
    "https://clinicaltrials.gov/study/NCT01276470",
    "https://clinicaltrials.gov/study/NCT00313066",
    "https://clinicaltrials.gov/study/NCT01239966",
    "https://clinicaltrials.gov/study/NCT01752049",
    "https://clinicaltrials.gov/study/NCT01066286",
    "https://clinicaltrials.gov/study/NCT01207518",
    "https://clinicaltrials.gov/study/NCT02067039",
    "https://clinicaltrials.gov/study/NCT00336076",
    "https://clinicaltrials.gov/study/NCT01692327",
    "https://clinicaltrials.gov/study/NCT01289275",
    "https://clinicaltrials.gov/study/NCT01282853",
    "https://clinicaltrials.gov/study/NCT00571389",
    "https://clinicaltrials.gov/study/NCT01806987",
    "https://clinicaltrials.gov/study/NCT02099435",
    "https://clinicaltrials.gov/study/NCT01804179",
    "https://clinicaltrials.gov/study/NCT01759901",
    "https://clinicaltrials.gov/study/NCT02086955",
    "https://clinicaltrials.gov/study/NCT01513980",
    "https://clinicaltrials.gov/study/NCT01688271",
    "https://clinicaltrials.gov/study/NCT00044304",
    "https://clinicaltrials.gov/study/NCT00266981",
    "https://clinicaltrials.gov/study/NCT01678859",
    "https://clinicaltrials.gov/study/NCT01361334",
    "https://clinicaltrials.gov/study/NCT01150279",
    "https://clinicaltrials.gov/study/NCT01874665",
    "https://clinicaltrials.gov/study/NCT01847911",
    "https://clinicaltrials.gov/study/NCT01833910",
    "https://clinicaltrials.gov/study/NCT01602939",
    "https://clinicaltrials.gov/study/NCT00744198",
    "https://clinicaltrials.gov/study/NCT01276951",
    "https://clinicaltrials.gov/study/NCT00874289",
    "https://clinicaltrials.gov/study/NCT00276926",
    "https://clinicaltrials.gov/study/NCT00608725",
    "https://clinicaltrials.gov/study/NCT01742065",
    "https://clinicaltrials.gov/study/NCT01882842",
    "https://clinicaltrials.gov/study/NCT01446120",
    "https://clinicaltrials.gov/study/NCT01058252",
    "https://clinicaltrials.gov/study/NCT01824615",
    "https://clinicaltrials.gov/study/NCT01656616",
    "https://clinicaltrials.gov/study/NCT00849407",
    "https://clinicaltrials.gov/study/NCT01532076",
    "https://clinicaltrials.gov/study/NCT01774266",
    "https://clinicaltrials.gov/study/NCT01941264",
    "https://clinicaltrials.gov/study/NCT01776736",
    "https://clinicaltrials.gov/study/NCT00743418",
    "https://clinicaltrials.gov/study/NCT01322698",
    "https://clinicaltrials.gov/study/NCT01978210",
    "https://clinicaltrials.gov/study/NCT01008228",
    "https://clinicaltrials.gov/study/NCT01498029",
    "https://clinicaltrials.gov/study/NCT01396148",
    "https://clinicaltrials.gov/study/NCT01340105",
    "https://clinicaltrials.gov/study/NCT01777529",
    "https://clinicaltrials.gov/study/NCT01532765",
    "https://clinicaltrials.gov/study/NCT02083042",
    "https://clinicaltrials.gov/study/NCT02156427",
    "https://clinicaltrials.gov/study/NCT01830361",
    "https://clinicaltrials.gov/study/NCT01806571",
    "https://clinicaltrials.gov/study/NCT00794651",
    "https://clinicaltrials.gov/study/NCT01852071",
    "https://clinicaltrials.gov/study/NCT01738139",
    "https://clinicaltrials.gov/study/NCT02171286",
    "https://clinicaltrials.gov/study/NCT02005861",
    "https://clinicaltrials.gov/study/NCT01219452",
    "https://clinicaltrials.gov/study/NCT01954212",
    "https://clinicaltrials.gov/study/NCT01559168",

]
links_KRAS = [
    "https://clinicaltrials.gov/study/NCT01871311",
    "https://clinicaltrials.gov/study/NCT01838577",
    "https://clinicaltrials.gov/study/NCT02015117",
    "https://clinicaltrials.gov/study/NCT01380795",
    "https://clinicaltrials.gov/study/NCT01190462",
    "https://clinicaltrials.gov/study/NCT00856375",
    "https://clinicaltrials.gov/study/NCT01508000",
    "https://clinicaltrials.gov/study/NCT02129257",
    "https://clinicaltrials.gov/study/NCT00842257",
    "https://clinicaltrials.gov/study/NCT02135757",
    "https://clinicaltrials.gov/study/NCT01294826",
    "https://clinicaltrials.gov/study/NCT01933932",
    "https://clinicaltrials.gov/study/NCT01384994",
    "https://clinicaltrials.gov/study/NCT00779454",
    "https://clinicaltrials.gov/study/NCT01693419",
    "https://clinicaltrials.gov/study/NCT01260415",
    "https://clinicaltrials.gov/study/NCT01320254",
    "https://clinicaltrials.gov/study/NCT01750918",
    "https://clinicaltrials.gov/study/NCT01935973",
    "https://clinicaltrials.gov/study/NCT00964457",
    "https://clinicaltrials.gov/study/NCT01740648",
    "https://clinicaltrials.gov/study/NCT01651013",
    "https://clinicaltrials.gov/study/NCT01752400",
    "https://clinicaltrials.gov/study/NCT02039336",
    "https://clinicaltrials.gov/study/NCT01449058",
    "https://clinicaltrials.gov/study/NCT01646554",
    "https://clinicaltrials.gov/study/NCT01358812",
    "https://clinicaltrials.gov/study/NCT01704703",
    "https://clinicaltrials.gov/study/NCT01740804",
    "https://clinicaltrials.gov/study/NCT01542437",
    "https://clinicaltrials.gov/study/NCT01836653",
    "https://clinicaltrials.gov/study/NCT01206049",
    "https://clinicaltrials.gov/study/NCT01124669",
    "https://clinicaltrials.gov/study/NCT01574300",
    "https://clinicaltrials.gov/study/NCT01282502",
    "https://clinicaltrials.gov/study/NCT01110785",
    "https://clinicaltrials.gov/study/NCT01787500",
    "https://clinicaltrials.gov/study/NCT01986166",
    "https://clinicaltrials.gov/study/NCT01596790",
    "https://clinicaltrials.gov/study/NCT01719380",
    "https://clinicaltrials.gov/study/NCT01394120",
    "https://clinicaltrials.gov/study/NCT01306045",
    "https://clinicaltrials.gov/study/NCT01943786",
    "https://clinicaltrials.gov/study/NCT01892527",
    "https://clinicaltrials.gov/study/NCT01360931",
    "https://clinicaltrials.gov/study/NCT01100840",
    "https://clinicaltrials.gov/study/NCT02041468",
    "https://clinicaltrials.gov/study/NCT01802645",
    "https://clinicaltrials.gov/study/NCT01688232",
    "https://clinicaltrials.gov/study/NCT01829217",
    "https://clinicaltrials.gov/study/NCT02171286",
    "https://clinicaltrials.gov/study/NCT01912625",

]

# Fetch data and create documents from the list of links
fetch_data_and_create_documents_from_links(links_ALK, "clinical_trials_custom_links_gemini", "clinical_trials", "ALK",
                                           "ALK")
fetch_data_and_create_documents_from_links(links_BRAF, "clinical_trials_custom_links_gemini", "clinical_trials",
                                           "BRAF", "BRAF")
fetch_data_and_create_documents_from_links(links_EGFR, "clinical_trials_custom_links_gemini", "clinical_trials",
                                           "EGFR", "EGFR")
fetch_data_and_create_documents_from_links(links_ERBB2, "clinical_trials_custom_links_gemini", "clinical_trials",
                                           "ERBB2", "ERBB2")
fetch_data_and_create_documents_from_links(links_KIT, "clinical_trials_custom_links_gemini", "clinical_trials", "KIT",
                                           "KIT")
fetch_data_and_create_documents_from_links(links_KRAS, "clinical_trials_custom_links_gemini", "clinical_trials",
                                           "KRAS", "KRAS")
